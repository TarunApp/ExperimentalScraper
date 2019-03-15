from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT


#Open Document
document = Document()
section = document.sections[0]

#FORMATTING - LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENTATION.LANDSCAPE 
section.page_width = new_width
section.page_height = new_height

#Styling
header = section.header
paragraphtitle = header.paragraphs[0]
paragraphtitle.text = "The Verona Times"
paragraphtitle.paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER




#Table Formats
table = document.add_table(rows=5, cols=5)

Monday = table.cell(0,2)
Mondayword = table.cell(1, 2)

Tuesday = table.cell(0,0)
Tuesdayword = table.cell(4,0)

Tuesday.text = 'test 1 Monday testing'
Tuesdayword.text = 'a;sldkjf;alsidfj'
Monday.text = '1908 – Butch Cassidy and the Sundance Kid are killed by Bolivian officials in a gunfight.'
Mondayword.text = 'Verminllion(n) - A vivid red to reddish orange color'


print("Done")
document.save('testfile.docx')